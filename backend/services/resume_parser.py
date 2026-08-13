"""Extract plain text from uploaded PDF/DOCX resumes."""

import io

import fitz  # PyMuPDF
from docx import Document
from fastapi import HTTPException

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB
ALLOWED_EXTENSIONS = (".pdf", ".docx")


def validate_file(filename: str, content: bytes) -> str:
    name = (filename or "").lower()
    if not name.endswith(ALLOWED_EXTENSIONS):
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Please upload a PDF or DOCX resume.",
        )
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail="File is too large. Maximum allowed size is 5 MB.",
        )
    if not content:
        raise HTTPException(status_code=400, detail="The uploaded file is empty.")
    return ".pdf" if name.endswith(".pdf") else ".docx"


def _extract_pdf(content: bytes) -> str:
    try:
        with fitz.open(stream=content, filetype="pdf") as doc:
            return "\n".join(page.get_text() for page in doc)
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=400, detail=f"Could not read the PDF: {exc}") from exc


def _extract_docx(content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=400, detail=f"Could not read the DOCX: {exc}") from exc


def extract_text(filename: str, content: bytes) -> str:
    ext = validate_file(filename, content)
    text = _extract_pdf(content) if ext == ".pdf" else _extract_docx(content)
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(text) < 50:
        raise HTTPException(
            status_code=400,
            detail="Could not extract readable text from this resume. Try a text-based PDF or a DOCX file.",
        )
    return text
